"""Assemble multiple RTF files into a single RTF or DOCX file."""

import os
from pathlib import Path
from typing import Union

from .input import RTFPage


def assemble_rtf(
    input_files: list[str],
    output_file: str,
    landscape: bool = False,
) -> None:
    """Combine multiple RTF files into a single RTF file.

    Args:
        input_files: List of paths to input RTF files.
        output_file: Path to the output RTF file.
        landscape: Whether the output should be landscape. Defaults to False.
            Note: This argument is currently not used in the RTF assembly logic
            as it preserves original file content, but kept for API consistency.
    """
    if not input_files:
        raise ValueError("Input files list cannot be empty")

    # Check input files exist
    missing_files = [f for f in input_files if not os.path.exists(f)]
    if missing_files:
        raise FileNotFoundError(f"Missing files: {', '.join(missing_files)}")

    # Read all RTF files
    rtf_contents = []
    for f in input_files:
        with open(f, "r", encoding="utf-8") as file:
            rtf_contents.append(file.readlines())

    if not rtf_contents:
        return

    # Process RTF content
    # 1. Keep the first file's header (up to fcharset/fonttbl)
    # 2. For all files, extract the body content
    # 3. Join with page breaks
    
    final_lines = []
    
    # Helper to find start index based on fcharset
    def find_start_index(lines):
        for i, line in enumerate(lines):
            if "fcharset" in line:
                return i + 2
        return 0

    new_page_cmd = r"\page" + "\n"

    processed_parts = []
    
    for i, lines in enumerate(rtf_contents):
        start_idx = 0
        if i > 0:
            # For subsequent files, skip header
            start_idx = find_start_index(lines)
            
        end_idx = len(lines)
        if i < len(rtf_contents) - 1:
            # Remove last line (closing brace) for all but last file
            if lines[-1].strip() == "}":
                end_idx -= 1
        
        part = lines[start_idx:end_idx]
        processed_parts.extend(part)
        
        if i < len(rtf_contents) - 1:
            processed_parts.append(new_page_cmd)

    # Write output
    with open(output_file, "w", encoding="utf-8") as f:
        f.writelines(processed_parts)


def assemble_docx(
    input_files: list[str],
    output_file: str,
    landscape: Union[bool, list[bool]] = False,
) -> None:
    """Combine multiple RTF files into a single DOCX file.

    Args:
        input_files: List of paths to input RTF files.
        output_file: Path to the output DOCX file.
        landscape: Whether the output should be landscape. Can be a single bool
            (applies to all) or a list of bools (one per file). Defaults to False.
    """
    try:
        import docx
        from docx.enum.section import WD_ORIENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches
    except ImportError as e:
        raise ImportError(
            "python-docx is required for assemble_docx. "
            "Install it with: pip install 'rtflite[docx]'"
        ) from e

    if not input_files:
        raise ValueError("Input files list cannot be empty")

    # Check input files exist
    missing_files = [f for f in input_files if not os.path.exists(f)]
    if missing_files:
        raise FileNotFoundError(f"Missing files: {', '.join(missing_files)}")

    # Handle landscape argument
    if isinstance(landscape, bool):
        landscape_list = [landscape] * len(input_files)
    else:
        if len(landscape) != len(input_files):
            raise ValueError("Length of landscape list must match input files")
        landscape_list = landscape

    # Create new document
    doc = docx.Document()
    
    for i, (input_file, is_landscape) in enumerate(zip(input_files, landscape_list)):
        # Absolute path needed for fields
        abs_path = os.path.abspath(input_file)
        
        # Escape backslashes for the field code
        path_str = abs_path.replace("\\", "\\\\")
        
        # Create INCLUDETEXT field
        field_code = f'INCLUDETEXT "{path_str}"'
        
        # Add "Table X" caption
        p = doc.add_paragraph()
        p.add_run("Table ")
        _add_field(p, r"SEQ Table \* ARABIC")
        p.add_run("\n") # Linebreak
        
        # Add the INCLUDETEXT field
        _add_field(p, field_code)
        
        # Handle section breaks and orientation
        if i < len(input_files) - 1:
            doc.add_section()
        
        # Set orientation for the current section
        section = doc.sections[-1]
        if is_landscape:
            section.orientation = WD_ORIENT.LANDSCAPE
            w, h = section.page_width, section.page_height
            if w < h: # If currently portrait
                section.page_width = h
                section.page_height = w
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            w, h = section.page_width, section.page_height
            if w > h: # If currently landscape
                section.page_width = h
                section.page_height = w

    doc.save(output_file)


def _add_field(paragraph, field_code):
    """Add a complex field to a paragraph."""
    # This is low-level XML manipulation for python-docx to add fields
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)

    run = paragraph.add_run()
    r = run._r
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_code
    r.append(instrText)

    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    r.append(fldChar)

    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)

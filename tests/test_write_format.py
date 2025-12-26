from html.parser import HTMLParser
from pathlib import Path

import polars as pl
import pytest

from rtflite.encode import RTFDocument
from rtflite.input import RTFBody, RTFColumnHeader, RTFTitle
from tests.conftest import (
    skip_if_no_libreoffice,
    skip_if_no_libreoffice_and_pypdf,
    skip_if_no_libreoffice_and_python_docx,
)


def _normalize_extracted_text(text: str) -> str:
    return " ".join(text.split())


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._chunks: list[str] = []
        self._ignored_tags: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        del attrs
        if tag in {"script", "style"}:
            self._ignored_tags.append(tag)

    def handle_endtag(self, tag: str) -> None:
        if self._ignored_tags and self._ignored_tags[-1] == tag:
            self._ignored_tags.pop()

    def handle_data(self, data: str) -> None:
        if self._ignored_tags:
            return
        self._chunks.append(data)

    def text(self) -> str:
        return " ".join(self._chunks)


def _html_to_text(html: str) -> str:
    parser = _HTMLTextExtractor()
    parser.feed(html)
    parser.close()
    return parser.text()


@pytest.fixture
def sample_document() -> RTFDocument:
    """Create a small RTFDocument for export integration tests."""
    df = pl.DataFrame({"A": ["alpha"], "B": ["beta"]})
    return RTFDocument(
        df=df,
        rtf_title=RTFTitle(text="Sample Title"),
        rtf_column_header=RTFColumnHeader(text=["A", "B"]),
        rtf_body=RTFBody(col_rel_width=[1, 1]),
    )


@pytest.mark.parametrize(
    ("method_name", "suffix"),
    [
        pytest.param(
            "write_docx",
            "docx",
            marks=skip_if_no_libreoffice_and_python_docx,
            id="docx",
        ),
        pytest.param(
            "write_html",
            "html",
            marks=skip_if_no_libreoffice,
            id="html",
        ),
        pytest.param(
            "write_pdf",
            "pdf",
            marks=skip_if_no_libreoffice_and_pypdf,
            id="pdf",
        ),
    ],
)
def test_write_export_creates_output_with_expected_content(
    sample_document: RTFDocument,
    tmp_path: Path,
    method_name: str,
    suffix: str,
):
    """Export writes file and preserves basic table content."""
    output_path = tmp_path / "reports" / f"table.{suffix}"
    getattr(sample_document, method_name)(output_path)

    assert output_path.exists()

    extracted_text: str
    if suffix == "docx":
        import docx

        document = docx.Document(str(output_path))
        extracted_text = _normalize_extracted_text(
            " ".join(
                [p.text for p in document.paragraphs]
                + [
                    cell.text
                    for table in document.tables
                    for row in table.rows
                    for cell in row.cells
                ]
            )
        )
    elif suffix == "html":
        html = output_path.read_text(encoding="utf-8", errors="ignore")
        extracted_text = _normalize_extracted_text(_html_to_text(html))
    else:
        from pypdf import PdfReader

        reader = PdfReader(str(output_path))
        extracted_text = _normalize_extracted_text(
            " ".join((page.extract_text() or "") for page in reader.pages)
        )

    assert "Sample Title" in extracted_text
    assert "alpha" in extracted_text
    assert "beta" in extracted_text

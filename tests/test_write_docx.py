import polars as pl
import pytest

from rtflite.encode import RTFDocument
from rtflite.input import RTFBody, RTFColumnHeader, RTFTitle
from tests.conftest import skip_if_no_libreoffice_and_python_docx

pytestmark = skip_if_no_libreoffice_and_python_docx


@pytest.fixture
def sample_document() -> RTFDocument:
    """Create a small RTFDocument for DOCX export tests."""
    df = pl.DataFrame({"A": ["alpha"], "B": ["beta"]})
    return RTFDocument(
        df=df,
        rtf_title=RTFTitle(text="Sample Title"),
        rtf_column_header=RTFColumnHeader(text=["A", "B"]),
        rtf_body=RTFBody(col_rel_width=[1, 1]),
    )


def test_write_docx_creates_docx(sample_document: RTFDocument, tmp_path):
    """DOCX export writes file and preserves table content."""
    output_path = tmp_path / "reports" / "table.docx"
    sample_document.write_docx(output_path)

    assert output_path.exists()

    import docx

    document = docx.Document(output_path)
    extracted_text = " ".join(
        [p.text for p in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )

    assert "Sample Title" in extracted_text
    assert "alpha" in extracted_text
    assert "beta" in extracted_text


def test_write_docx_uses_temp_files(sample_document: RTFDocument, tmp_path):
    """DOCX export should not leave intermediate artifacts in output dir."""
    output_dir = tmp_path / "final"
    output_path = output_dir / "clean.docx"

    sample_document.write_docx(output_path)

    assert output_path.exists()
    assert not any(path.suffix == ".rtf" for path in output_dir.iterdir())

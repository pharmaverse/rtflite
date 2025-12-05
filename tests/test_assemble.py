import pytest

from rtflite import assemble_docx, assemble_rtf, concatenate_docx
from tests.skip_conditions import skip_if_no_python_docx


@pytest.fixture
def sample_rtf_files(tmp_path):
    """Create sample RTF files for testing."""
    file1 = tmp_path / "test1.rtf"
    file2 = tmp_path / "test2.rtf"

    content1 = r"""{\rtf1\ansi\deff0
{\fonttbl{\f0 Arial;}}
{\colortbl;\red0\green0\blue0;}
\pard\plain\f0\fs24
Content of file 1
\par
}"""

    content2 = r"""{\rtf1\ansi\deff0
{\fonttbl{\f0 Arial;}}
{\colortbl;\red0\green0\blue0;}
\pard\plain\f0\fs24
Content of file 2
\par
}"""

    file1.write_text(content1, encoding="utf-8")
    file2.write_text(content2, encoding="utf-8")

    return [str(file1), str(file2)]


@pytest.fixture
def complex_rtf_files(tmp_path):
    """Create RTF files with complex headers."""
    file1 = tmp_path / "complex1.rtf"
    file2 = tmp_path / "complex2.rtf"

    header = r"""{\rtf1\ansi\deff0
{\fonttbl{\f0\froman\fcharset1\fprq2 Times New Roman;}
{\f1\froman\fcharset161\fprq2 Times New Roman Greek;}
}
"""
    content1 = header + r"\pard Content 1\par}"
    content2 = header + r"\pard Content 2\par}"

    file1.write_text(content1, encoding="utf-8")
    file2.write_text(content2, encoding="utf-8")

    return [str(file1), str(file2)]


@pytest.fixture
def sample_docx_files(tmp_path):
    """Create DOCX files for testing DOCX concatenation."""
    import docx
    from docx.enum.section import WD_ORIENT

    portrait_path = tmp_path / "portrait.docx"
    landscape_path = tmp_path / "landscape.docx"

    portrait_doc = docx.Document()
    portrait_doc.add_paragraph("Portrait content")
    portrait_doc.save(portrait_path)

    landscape_doc = docx.Document()
    landscape_doc.add_paragraph("Landscape content")
    section = landscape_doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    width, height = section.page_width, section.page_height
    if width is not None and height is not None and width < height:
        section.page_width, section.page_height = height, width
    landscape_doc.save(landscape_path)

    return [portrait_path, landscape_path]


def test_assemble_rtf_complex(complex_rtf_files, tmp_path):
    output_file = tmp_path / "combined_complex.rtf"
    assemble_rtf(complex_rtf_files, str(output_file))

    assert output_file.exists()
    content = output_file.read_text(encoding="utf-8")

    assert "Content 1" in content
    assert "Content 2" in content
    # Ensure header from second file is stripped (no double fonttbl)
    # We expect only one fonttbl block
    assert content.count(r"{\fonttbl") == 1


def test_assemble_rtf(sample_rtf_files, tmp_path):
    output_file = tmp_path / "combined.rtf"
    assemble_rtf(sample_rtf_files, str(output_file))

    assert output_file.exists()
    content = output_file.read_text(encoding="utf-8")

    # Check if content from both files is present
    assert "Content of file 1" in content
    assert "Content of file 2" in content

    # Check for page break
    assert r"\page" in content


def test_assemble_rtf_missing_file():
    with pytest.raises(FileNotFoundError):
        assemble_rtf(["non_existent.rtf"], "output.rtf")


def test_assemble_rtf_empty_list(tmp_path):
    output_file = tmp_path / "output.rtf"
    assemble_rtf([], str(output_file))
    assert not output_file.exists()


@skip_if_no_python_docx
def test_assemble_docx(sample_rtf_files, tmp_path):
    import docx

    output_file = tmp_path / "combined.docx"
    assemble_docx(sample_rtf_files, str(output_file))

    assert output_file.exists()

    doc = docx.Document(str(output_file))
    assert len(doc.paragraphs) > 0
    assert any("Table " in p.text for p in doc.paragraphs)


@skip_if_no_python_docx
def test_assemble_docx_landscape(sample_rtf_files, tmp_path):
    import docx
    from docx.enum.section import WD_ORIENT

    output_file = tmp_path / "combined_landscape.docx"
    assemble_docx(sample_rtf_files, str(output_file), landscape=[False, True])

    assert output_file.exists()

    doc = docx.Document(str(output_file))
    assert len(doc.sections) == 2
    assert doc.sections[0].orientation == WD_ORIENT.PORTRAIT
    assert doc.sections[1].orientation == WD_ORIENT.LANDSCAPE


@skip_if_no_python_docx
def test_assemble_docx_missing_file():
    with pytest.raises(FileNotFoundError):
        assemble_docx(["non_existent.rtf"], "output.docx")


@skip_if_no_python_docx
def test_concatenate_docx(sample_docx_files, tmp_path):
    import docx
    from docx.enum.section import WD_ORIENT

    output_file = tmp_path / "combined-docx.docx"
    concatenate_docx(
        [str(path) for path in sample_docx_files],
        output_file,
        landscape=[False, True],
    )

    assert output_file.exists()

    doc = docx.Document(str(output_file))
    assert len(doc.sections) == 2
    assert doc.sections[0].orientation == WD_ORIENT.PORTRAIT
    assert doc.sections[1].orientation == WD_ORIENT.LANDSCAPE

    combined_text = " ".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Portrait content" in combined_text
    assert "Landscape content" in combined_text


@skip_if_no_python_docx
def test_concatenate_docx_validates_landscape_length(sample_docx_files, tmp_path):
    output_file = tmp_path / "should-not-exist.docx"
    with pytest.raises(ValueError):
        concatenate_docx(
            [str(path) for path in sample_docx_files],
            output_file,
            landscape=[True],
        )


@skip_if_no_python_docx
def test_concatenate_docx_missing_file(tmp_path):
    with pytest.raises(FileNotFoundError):
        concatenate_docx([tmp_path / "missing.docx"], tmp_path / "out.docx")
